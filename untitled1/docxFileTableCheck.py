# -*- coding:UTF-8 -*-
import sys


import xml.etree.ElementTree as ET
import datetime
import xml.etree.ElementTree as XETree
from docx import Document
from openpyxl import load_workbook
sourceFilePath = r'C:\\PyCharm\\untitled1\\111\\1.docx'
doc = Document(sourceFilePath)
tables = doc.tables



def printTableContent(tables, outerIndex=-1):
    row_Title =[]
    row_xml = []
    for tbIndex, table in enumerate(tables):
        print(' ')
        if outerIndex != -1:
            print('=========================table', outerIndex, '   ========>cellTable:', tbIndex)
        else:
            print('=========================table', tbIndex)
        xml =table._tblPr.xml
        root_elem = ET.fromstring(xml)
        for ch in root_elem:
            key = ch.tag.split('}')[1]
            if key == 'tblCaption':
                titleName = str(list(ch.attrib.values())).split('\'')[1]
                print(key+":__"+titleName)
        for rIndex, row in enumerate(table.rows):
            row_content = []
            for  cell in row.cells:
                cText = cell.text
                row_content.append(cText)
                if len(cell.tables) > 0:
                    cIndex = 0
                    isTrue = 0
                    while cIndex < len(cell.tables):
                        rXml = cell.tables[cIndex]._tblPr.xml
                        cIndex +=1
                        if rXml not in row_xml:
                            isTrue = 1
                            row_xml.append(rXml)
                            continue
                    if isTrue == 1:
                        printTableContent(cell.tables, tbIndex)


            #print('--row', rIndex, '  ', row_content)

now = datetime.datetime.now()
print('1: ' + str(now))
printTableContent(tables)
now = datetime.datetime.now()
print('2: ' + str(now))