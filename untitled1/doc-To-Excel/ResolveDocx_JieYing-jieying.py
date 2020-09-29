# _*_ coding:utf-8 _*_
import os
import os.path
import sys
import xml.etree.ElementTree as XETree
import xml.etree.ElementTree as ET
from docx import Document
from openpyxl import load_workbook

cdfp = None
cwb = None
clws = None
DATANOTFOUND = 0
writeLog = 0
def arearMapSupInfoToExcelName(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols, reportType, dataSource, tableCode, dcolsNames):
    sheet = cwb[dSheet]
    sList = sCols.split(',')
    dList = dCols.split(',')
    dNameList = dcolsNames.split(',')
    dRowIncress = 0
    itemName = 1
    while sBeginRow <= sEndRow:
        for sCellI in range(len(sList)):
            sCol = sList[sCellI]
            value = table.rows[sBeginRow].cells[int(sCol)].text
            dRow = dBeginRow + dRowIncress
            sheet[dList[0] + str(dRow)] = dNameList[sCellI]
            sheet[dList[1] + str(dRow)] = itemName
            sheet[dList[2] + str(dRow)] = value
            sheet[dList[3] + str(dRow)] = reportType
            sheet[dList[4] + str(dRow)] = dataSource
            sheet[dList[5] + str(dRow)] = tableCode
            dRowIncress += 1
        sBeginRow += 1
        itemName += 1

def arearMapSupInfoToExcel(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols, reportType, dataSource, tableCode):
    sheet = cwb[dSheet]
    sList = sCols.split(',')
    dList = dCols.split(',')
    dRowIncress = 0
    itemName = 1
    while sBeginRow <= sEndRow:
        for sCellI in range(len(sList)):
            sCol = sList[sCellI]
            value = table.rows[sBeginRow].cells[int(sCol)].text
            dRow = dBeginRow + dRowIncress
            sheet[dList[0] + str(dRow)] = 'ItemCode'+ str(sCellI)
            sheet[dList[1] + str(dRow)] = itemName
            sheet[dList[2] + str(dRow)] = value
            sheet[dList[3] + str(dRow)] = reportType
            sheet[dList[4] + str(dRow)] = dataSource
            sheet[dList[5] + str(dRow)] = tableCode
            dRowIncress += 1
        sBeginRow += 1
        itemName += 1

def arearMapExtractDataToExcel(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols):
    sheet = cwb[dSheet]
    sList = sCols.split(',')
    dList = dCols.split(',')
    dRowIncress = 0
    while sBeginRow <= sEndRow:
        for sCellI in range(len(sList)):
            sCol = sList[sCellI]
            value = table.rows[sBeginRow].cells[int(sCol)].text
            dRow = dBeginRow + dRowIncress
            if dSheet == '资产统计信息':
                sheet['A' + str(dRow)] =  'ItemCode'+ str(dRowIncress)
            if value != '':
                sheet[dList[sCellI] + str(dRow)] = value
        dRowIncress += 1
        sBeginRow += 1

def cellMapExtractDataToExcel(table, dNode, dSheet):
    sheet = cwb[dSheet]
    for cell in dNode:
        cUsing = cell.attrib['using']
        cTag = cell.tag
        cText = cell.text
        if cUsing == 'replace':
            r = int(cText.split(',')[0])
            c = int(cText.split(',')[1])
            v = table.rows[r].cells[c].text.strip()
            if v != '':
                sheet[cTag] = v
        elif cUsing == 'sum':
            dcs = cText.split(';')
            sumv = 0
            for i in range(len(dcs)):
                r = int(dcs[i].split(',')[0])
                c = int(dcs[i].split(',')[1])
                v = table.rows[r].cells[c].text.strip()
                v = v.replace(',', '').replace('-', '')
                if v != '':
                    sumv += float(v)
            sheet[cTag] = "{0:.2f}".format(sumv)

def arearMapExtract(table, cfgItem, itemIndex):
    global DATANOTFOUND
    itemDesc = cfgItem.attrib['desc']
    sNode = cfgItem.find('source')
    dNode = cfgItem.find('dest')
    sAnchor = sNode.attrib['anchor'].strip()
    sSkipRows = int(sNode.attrib['skiprows']) if 'skiprows' in sNode.attrib else 0
    sAnchorEnd = sNode.attrib['anchorend'].strip()
    dLimit = int(dNode.attrib['limited']) if 'limited' in dNode.attrib else 0
    sAnchorEndArr = sAnchorEnd.split('$')
    sBeginRow = -1
    sEndRow = -1
    for rIndex, row in enumerate(table.rows):
        firstCellText = row.cells[0].text.strip()
        if firstCellText == '':
            continue
        if sBeginRow == -1 and (firstCellText.startswith(sAnchor) or firstCellText.endswith(sAnchor)):
            sBeginRow = rIndex + sSkipRows + 1;
        elif sBeginRow != -1 and sAnchorEnd != '' and (
                (sAnchorEnd.find('$') == -1 and firstCellText.startswith(sAnchorEnd)) or (
                sAnchorEnd.find('$') != -1 and firstCellText in sAnchorEndArr)):
            sEndRow = rIndex if dLimit == 0 or rIndex + 1 - sBeginRow <= dLimit else sBeginRow + dLimit - 1
            break
        if sBeginRow != -1 and sEndRow == -1:
            rowsCount = len(table.rows)
            if dLimit == 0 and sAnchorEnd == '':
                sEndRow = rowsCount - 1
                break
            if dLimit != 0 and sAnchorEnd == '':
                sEndRow = sBeginRow + dLimit if sBeginRow + dLimit <= rowsCount - 1 else rowsCount - 1
                break
            if dLimit != 0 and sAnchorEnd != '' and rIndex - sBeginRow == dLimit - 1:
                sEndRow = rIndex
                break
    if sBeginRow != -1 and sEndRow != -1:
        sCols = sNode.attrib['cols']
        dCols = dNode.attrib['cols']
        dSheet = dNode.attrib['sheet']
        dBeginRow = int(dNode.attrib['beginrow'])
        writeSheetLog('{0} 提取： 【{1}】'.format(itemIndex + 1, itemDesc))
        writeSheetLog(
            '--------源表格起始行:{0},源表格结束行:{1},目标Sheet[{3}]开始行:{2}'.format(sBeginRow, sEndRow, dBeginRow, dSheet))
        if 'type' in cfgItem.attrib:
            reportType = dNode.attrib['ReportType']
            dataSource = dNode.attrib['DataSource']
            tableCode = dNode.attrib['TableCode']
            if 'colsNames' in sNode.attrib:
                dcolsNames = sNode.attrib['colsNames']
                arearMapSupInfoToExcelName(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols, reportType, dataSource, tableCode, dcolsNames)
            else:
                arearMapSupInfoToExcel(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols, reportType, dataSource, tableCode)
        else:
            arearMapExtractDataToExcel(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols)
        writeSheetLog('--------【{0}】数据已提取完成'.format(itemDesc))
        if writeLog == 0:
            cwb.save(cdfp)
    if sBeginRow == -1 and sEndRow == -1:
        DATANOTFOUND += 1
        writeSheetLog('{1} 【{0}】数据未找到，请检查源文件和配置文件'.format(itemDesc, itemIndex))

def arearMapExtractTable(tables, cfgItem, itemIndex):
    global DATANOTFOUND
    itemDesc = cfgItem.attrib['desc']
    sNode = cfgItem.find('source')
    dNode = cfgItem.find('dest')
    sAnchor = sNode.attrib['anchor'].strip()
    sSkipRows = int(sNode.attrib['skiprows']) if 'skiprows' in sNode.attrib else 0
    sAnchorEnd = sNode.attrib['anchorend'].strip()
    dLimit = int(dNode.attrib['limited']) if 'limited' in dNode.attrib else 0
    sAnchorEndArr = sAnchorEnd.split('$')
    sBeginRow = -1
    sEndRow = -1
    index = int(sNode.attrib['index'].strip())
    for tbIndex, table in enumerate(tables):
        if tbIndex >= index or index == -1:
            for rIndex, row in enumerate(table.rows):
                firstCellText = row.cells[0].text.strip()
                if firstCellText == '':
                    continue
                if sBeginRow == -1 and (firstCellText.startswith(sAnchor) or firstCellText.endswith(sAnchor)):
                    sBeginRow = rIndex + sSkipRows + 1;
                elif sBeginRow != -1 and sAnchorEnd != '' and (
                        (sAnchorEnd.find('$') == -1 and firstCellText.startswith(sAnchorEnd)) or (
                        sAnchorEnd.find('$') != -1 and firstCellText in sAnchorEndArr)):
                    sEndRow = rIndex if dLimit == 0 or rIndex + 1 - sBeginRow <= dLimit else sBeginRow + dLimit - 1
                    break
                if sBeginRow != -1 and sEndRow == -1:
                    rowsCount = len(table.rows)
                    if dLimit == 0 and sAnchorEnd == '':
                        sEndRow = rowsCount - 1
                        break
                    if dLimit != 0 and sAnchorEnd == '':
                        sEndRow = sBeginRow + dLimit if sBeginRow + dLimit <= rowsCount - 1 else rowsCount - 1
                        break
                    if dLimit != 0 and sAnchorEnd != '' and rIndex - sBeginRow == dLimit - 1:
                        sEndRow = rIndex
                        break
            if sBeginRow != -1 and sEndRow != -1:
                sCols = sNode.attrib['cols']
                dCols = dNode.attrib['cols']
                dSheet = dNode.attrib['sheet']
                dBeginRow = int(dNode.attrib['beginrow'])
                writeSheetLog('{0} 提取： 【{1}】'.format(itemIndex + 1, itemDesc))
                writeSheetLog(
                    '--------源表格起始行:{0},源表格结束行:{1},目标Sheet[{3}]开始行:{2}'.format(sBeginRow, sEndRow, dBeginRow, dSheet))
                if 'type' in cfgItem.attrib:
                    reportType = dNode.attrib['ReportType']
                    dataSource = dNode.attrib['DataSource']
                    tableCode = dNode.attrib['TableCode']
                    if 'colsNames' in sNode.attrib:
                        dcolsNames = sNode.attrib['colsNames']
                        arearMapSupInfoToExcelName(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols, reportType, dataSource, tableCode, dcolsNames)
                    else:
                        arearMapSupInfoToExcel(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols, reportType, dataSource, tableCode)
                else:
                    arearMapExtractDataToExcel(table, sBeginRow, sEndRow, sCols, dSheet, dBeginRow, dCols)
                writeSheetLog('--------【{0}】数据已提取完成'.format(itemDesc))
                if writeLog == 0:
                    cwb.save(cdfp)
                break
    if sBeginRow == -1 and sEndRow == -1:
        DATANOTFOUND += 1
        writeSheetLog('{1} 【{0}】数据未找到，请检查源文件和配置文件'.format(itemDesc, itemIndex))

def cellMapExtract(tables, cfgItem, itemIndex):
    global DATANOTFOUND
    itemDesc = cfgItem.attrib['desc']
    sNode = cfgItem.find('source')
    dNode = cfgItem.find('dest')
    foundTable = 0
    sAnchor = sNode.attrib['anchor'].strip()
    index = int(sNode.attrib['index'].strip())
    for tbIndex, table in enumerate(tables):
        if tbIndex >= index or index == -1:
            for rIndex, row in enumerate(table.rows):
                firstCellText = row.cells[0].text.strip()
                if firstCellText == '' or firstCellText != sAnchor:
                    continue
                if firstCellText == sAnchor:
                    foundTable = 1
                    break
            if foundTable == 1:
                dSheet = dNode.attrib['sheet']
                writeSheetLog('{0} 提取： 【{1}】'.format(itemIndex + 1, itemDesc))
                writeSheetLog('--------开始表格映射映射数据提取')
                cellMapExtractDataToExcel(table, dNode, dSheet)
                writeSheetLog('--------【{0}】数据已提取完成'.format(itemDesc))
                if writeLog == 0:
                    cwb.save(cdfp)
                break
    if foundTable == 0:
        DATANOTFOUND += 1
        writeSheetLog('\033[1;31m {1} 【{0}】数据未找到，请检查源文件和配置文件 \033[0m!'.format(itemDesc, itemIndex + 1))

def extractDocFile(cfgItems, sourceFilePath):
    doc = Document(sourceFilePath)
    tables = doc.tables
    for i in range(len(cfgItems)):
        cfgItem = cfgItems[i]
        if 'useTableName' not in cfgItem.attrib:
            if 'type' in cfgItem.attrib:
                itemType = cfgItem.attrib['type']
                if itemType == 'cellmap':
                    cellMapExtract(tables, cfgItem, i)
                elif itemType == 'supInfo':
                    arearMapExtractTable(tables, cfgItem, i)
            else:
                arearMapExtractTable(tables, cfgItem, i)
        else:
            if 'type' in cfgItem.attrib:
                itemType = cfgItem.attrib['type']
                if itemType == 'cellmap':
                    findTableName(tables, cfgItem, i, 1)
                elif itemType == 'supInfo':
                    findTableName(tables, cfgItem, i, 0)
            else:
                findTableName(tables, cfgItem, i, 0)

def writeSheetLog(info):
    if writeLog == 1 and clws is not None:
        clws['A' + str(clws.max_row + 1)] = info
        cwb.save(cdfp)

def findTableName(tables, cfgItem, index, typeId):
    sNode = cfgItem.find('source')
    dNode = cfgItem.find('dest')
    itemDesc = cfgItem.attrib['desc']
    tableName = sNode.attrib['tableName'].strip()
    for tbIndex, table in enumerate(tables):
        xml = table._tblPr.xml
        root_elem = ET.fromstring(xml)
        for ch in root_elem:
            key = ch.tag.split('}')[1]
            if key == 'tblCaption':
                titleName = str(list(ch.attrib.values())).split('\'')[1]
                if titleName == tableName:
                    if typeId == 1:
                        dSheet = dNode.attrib['sheet']
                        writeSheetLog('{0} 提取： 【{1}】'.format(tbIndex + 1, itemDesc))
                        writeSheetLog('--------开始表格映射映射数据提取')
                        cellMapExtractDataToExcel(table, dNode, dSheet)
                        writeSheetLog('--------【{0}】数据已提取完成'.format(itemDesc))
                        if writeLog == 0:
                            cwb.save(cdfp)
                        break
                    elif typeId == 0:
                        arearMapExtract(table, cfgItem, index)
        for rIndex, row in enumerate(table.rows):
            for cell in row.cells:
                if len(cell.tables) > 0:
                    findTableName(cell.tables, cfgItem,index, typeId)

def main():
    global DATANOTFOUND
    global cdfp
    global clws
    global cwb
    global writeLog
    reload(sys)
    sys.setdefaultencoding('utf-8')
    sourceFilePath = sys.argv[1]
    destFileName =sys.argv[3]
    configFilePath = sys.argv[2]

    mappingTree = XETree.parse(configFilePath)
    cfgRoot = mappingTree.getroot()
    destFolder = cfgRoot.attrib['destfolder']
    templateFilePath = cfgRoot.attrib['template']
    writeLog = int(cfgRoot.attrib['writelog']) if 'writelog' in cfgRoot.attrib else 0

    cdfp = os.path.join(destFolder, destFileName)
    if not os.path.exists(destFolder):
        os.makedirs(destFolder)
    if os.path.exists(cdfp):
        os.remove(cdfp)
    open(cdfp, "wb").write(open(templateFilePath, "rb").read())
    cwb = load_workbook(cdfp)
    if writeLog == 1:
        clws = cwb.create_sheet("Extract Log")
        cwb.save(cdfp)
    extractDocFile(cfgRoot, sourceFilePath)



main()
